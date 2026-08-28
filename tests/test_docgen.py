# -*- coding: utf-8 -*-
"""Word 文档生成 + 全流程管线测试。

用种子项目「个人网银系统」走完整第二批链路(OSV 用 MockTransport 模拟),
对生成的 4 份 Word 做内容级断言: 章节存在、关键数据可见、高危标红、密码策略
参数与引擎口径一致。离线降级路径单独覆盖。
"""
from datetime import datetime

import httpx
import pytest
from docx import Document
from docx.shared import RGBColor

import shared.constants as C
from models import VulnerabilityRecord, init_db, make_engine
from rules import RuleEngine
from rules.context import RequirementContext
from services.docgen import generate_all_documents, load_doc_style
from services.pipeline import _load_vulnerabilities, run_full_pipeline
from services.seed_data import seed_demo_project

VULN_RED = RGBColor.from_string("C00000")


def make_mock_client(payload_by_purl):
    """构造 MockTransport 版 OSV 客户端; 未登记 purl 返回空结果(官方无命中形态)。"""
    import json as _json

    calls: list[str] = []

    def handler(request: httpx.Request) -> httpx.Response:
        from services.osv import OsvClient

        purl = _json.loads(request.content)["package"]["purl"]
        calls.append(purl)
        payload = payload_by_purl.get(purl)
        if isinstance(payload, Exception):
            raise payload
        return httpx.Response(200, json=payload or {})

    from services.osv import OsvClient
    return OsvClient(transport=httpx.MockTransport(handler)), calls


LOG4J_PURL = "pkg:maven/org.apache.logging.log4j/log4j-core@2.14.1"
FASTJSON_PURL = "pkg:maven/com.alibaba/fastjson@1.2.70"

GHSA_LOG4J = {
    "id": "GHSA-jfh8-c2jp-5v3q",
    "summary": "Apache Log4j2 JNDI注入远程代码执行漏洞(Log4Shell)",
    "aliases": ["CVE-2021-44228"],
    "severity": [{"type": "CVSS_SCORE", "score": "10.0"}],
    "affected": [{
        "ranges": [{"events": [{"introduced": "2.0"}, {"fixed": "2.15.0"},
                               {"introduced": "2.15.1"}, {"fixed": "2.17.0"}]}],
    }],
    "database_specific": {"severity": "CRITICAL"},
}
CVE_FASTJSON = {
    "id": "GHSA-fastjson-test",
    "aliases": ["CVE-2022-25845"],
    "summary": "fastjson autoType未完全闭合导致远程代码执行",
    "affected": [{
        "ranges": [{"events": [{"introduced": "0"}, {"fixed": "2.0.0"}]}],
    }],
    "database_specific": {"severity": "HIGH"},
}


@pytest.fixture(scope="module")
def batch2_output(tmp_path_factory):
    """种子项目 → OSV同步(mock) → 规则引擎 → 4份文档, 全模块共用。"""
    engine = make_engine("sqlite:///:memory:")
    init_db(engine)
    from sqlalchemy.orm import sessionmaker
    session = sessionmaker(bind=engine, autoflush=False, expire_on_commit=False)()

    project = seed_demo_project(session)
    client, calls = make_mock_client({
        LOG4J_PURL: {"vulns": [GHSA_LOG4J]},
        FASTJSON_PURL: {"vulns": [dict(CVE_FASTJSON)]},
    })

    ctx = RequirementContext.from_db(session, project.id)
    from services.osv import sync_vulnerabilities
    _, sync_result = sync_vulnerabilities(
        session, ctx.components, client=client,
        now=datetime(2026, 8, 27, 10, 0, 0),
    )
    ctx = RequirementContext.from_db(session, project.id)

    reqs = RuleEngine.load().generate_and_save(ctx, session)
    vulns = _load_vulnerabilities(session, ctx.components)

    out_dir = tmp_path_factory.mktemp("docs")
    documents = generate_all_documents(
        ctx, out_dir, requirements=reqs, vulnerabilities=vulns,
        osv_summary=sync_result.summary_text(),
        generated_at=datetime(2026, 8, 27, 11, 30, 0),
    )
    yield session, project, reqs, vulns, sync_result, documents
    session.close()


# ────────────────────────── 文档读取辅助 ──────────────────────────


def iter_document_text(document: Document):
    """段落文本 + 表格单元格文本(含多段) 依序产出。"""
    for para in document.paragraphs:
        yield para.text
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para.text


def colored_red_texts(document: Document) -> set[str]:
    """收集标红字体的运行片段(紧急需求行/严重漏洞行)。"""
    texts = set()

    def scan(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                if _is_red(run):
                    texts.add(run.text)

    scan(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                scan(cell.paragraphs)
    return texts


def _is_red(run) -> bool:
    try:
        rgb = run.font.color.rgb
    except Exception:
        return False
    return rgb == VULN_RED


def full_text(document: Document) -> str:
    return "\n".join(iter_document_text(document))


# ────────────────────────── 断言 ──────────────────────────


def test_five_documents_generated(batch2_output):
    *_, documents = batch2_output
    assert set(documents) == {"grading", "requirement", "design", "sbom_vuln", "review"}
    for path in documents.values():
        assert path.exists() and path.stat().st_size > 3_000, f"{path} 过小或缺失"


def test_grading_report_sections_and_level(batch2_output):
    *_, documents = batch2_output
    text = full_text(Document(str(documents["grading"])))
    assert "一、项目基本信息" in text and "二、等保定级问卷答案" in text
    assert "三、定级结论" in text and "四、判定理由" in text and "六、人工修正栏" in text
    assert "五、判定依据" in text and "七、安全中心复核意见" in text
    assert "等保三级" in text and "「三级」" in text          # 种子建议定级为三级
    assert "处理敏感个人信息并涉及资金交易" in text            # 判定理由全文带入
    assert "最终定级：三级" in text                            # 种子问卷已人工确认(立项门禁前置)
    assert "项目编码：PRJ-IBANK-2026" in text                 # 封面元信息


def test_requirement_spec_groups_by_asvs_and_has_matrix_appendix(batch2_output):
    session, project, reqs, *_ , documents = batch2_output
    text = full_text(Document(str(documents["requirement"])))
    assert "一、需求概览" in text
    assert "附录A 权限矩阵" in text
    assert "交易流水记录" in text                       # 权限矩阵资源列
    assert "需走审批流程" in text                        # 图例说明
    assert "SEC-V12-001" in text                       # 上传功能需求在列

    critical_ids = {r.req_id for r in reqs if r.priority == "critical"}
    red = colored_red_texts(Document(str(documents["requirement"])))
    assert critical_ids & red, "紧急级需求的编号行应标红"


def test_design_baseline_contains_dictionary_policy_auth(batch2_output):
    *_, documents = batch2_output
    text = full_text(Document(str(documents["design"])))
    # 六个章节标题
    for head in ("一、软件/框架版本清单", "二、数据字典", "三、API 接口安全属性表",
                 "四、基础设施资产清单", "五、登录与密码策略设计说明", "六、认证方式设计说明"):
        assert head in text
    # 数据字典三级结构落到表里
    assert "指纹生物特征" in text and "t_biometric_template" in text \
        and "fingerprint_feature" in text
    # 密码策略参数化描述与 Step6 配置一致(10位/4类/60天)
    assert "口令最小长度不少于10位" in text and "复杂度须至少包含4类字符" in text
    assert "有效期60天" in text
    # 认证方式设计说明逐项展开
    assert "短信验证码：" in text and "第三方OAuth：" in text


def test_sbom_vuln_report_marks_high_risk_red_and_links_fixing_reqs(batch2_output):
    session, project, reqs, vulns, sync_result, documents = batch2_output
    document = Document(str(documents["sbom_vuln"]))
    text = full_text(document)
    tokens = set(text.split())

    assert "CycloneDX 1.5" in text and "log4j-core" in text
    assert {"CVE-2021-44228", "CVE-2022-25845"} <= tokens, "两枚演示 CVE 应出现在漏洞表中"

    red = colored_red_texts(document)
    assert "CVE-2021-44228" in red and "严重" in red     # 高危置顶标红
    # 整改联动需求: SEC-V14-801 模板按组件实例化
    vuln_req_ids = {
        r.req_id for r in reqs if r.source_entity_type == "sbom_component"
    }
    assert len(vuln_req_ids) == 2
    assert vuln_req_ids <= set(text.split())
    assert any("整改" in r.title for r in reqs if r.req_id in vuln_req_ids)


def test_offline_variant_documents_render_placeholder(tmp_path):
    """漏洞查询不可用时: SBOM清单照常交付, 漏洞节显示暂不可用提示。"""
    engine = make_engine("sqlite:///:memory:")
    init_db(engine)
    from sqlalchemy.orm import sessionmaker
    session = sessionmaker(bind=engine, autoflush=False, expire_on_commit=False)()
    project = seed_demo_project(session)
    ctx = RequirementContext.from_db(session, project.id)
    reqs = RuleEngine.load().generate_and_save(ctx, session)

    documents = generate_all_documents(
        ctx, tmp_path, requirements=reqs, vulnerabilities=[],
        osv_summary="OSV查询: 查询失败1(log4j-core)",
    )
    text = full_text(Document(str(documents["sbom_vuln"])))
    assert "本次漏洞查询暂不可用" in text
    assert "当前无高危及以上漏洞, 未触发漏洞整改类安全需求。" in text
    session.close()


def test_style_template_override_merges_partially(tmp_path):
    """版式模板文件部分覆盖: 只改一处颜色也不丢默认值。"""
    style_file = tmp_path / "doc_style.yml"
    style_file.write_text(
        "colors:\n  vuln_red: \"FF0000\"\n", encoding="utf-8"
    )
    style = load_doc_style(style_file)
    assert style["colors"]["vuln_red"] == "FF0000"
    assert style["fonts"]["body_cn"] == "宋体"   # 其余沿用默认
    assert load_doc_style()["colors"]["vuln_red"] == "C00000"


# ────────────────────────── 管线编排 ──────────────────────────


def test_run_full_pipeline_offline_generates_everything(tmp_path):
    """skip_osv 模式: 需求/SBOM/文档全链路可用且落库数量一致。"""
    engine = make_engine("sqlite:///:memory:")
    init_db(engine)
    from sqlalchemy.orm import sessionmaker
    session = sessionmaker(bind=engine, autoflush=False, expire_on_commit=False)()
    project = seed_demo_project(session)

    result = run_full_pipeline(
        session, project.id, out_dir=tmp_path / "out", skip_osv=True,
    )

    assert result.sync is None
    expected = RuleEngine.load().generate(
        RequirementContext.from_db(session, project.id)
    )
    assert len(result.requirements) == len(expected)
    assert not [r for r in result.requirements if r.category == "第三方组件风险"], \
        "离线模式无漏洞输入, 不应有联动需求"

    import json as _json
    bom = _json.loads(result.bom_path.read_text(encoding="utf-8"))
    assert bom["specVersion"] == "1.5" and len(bom["components"]) == 10
    assert set(result.documents) == {"grading", "requirement", "design", "sbom_vuln", "review"}
    for path in result.documents.values():
        assert path.exists()
    session.close()


def test_run_full_pipeline_with_mocked_osv_creates_vuln_requirements(tmp_path):
    """在线(被mock)模式: OSV命中 → 引擎追加第三方组件风险需求 + 文档联动表格。"""
    engine = make_engine("sqlite:///:memory:")
    init_db(engine)
    from sqlalchemy.orm import sessionmaker
    session = sessionmaker(bind=engine, autoflush=False, expire_on_commit=False)()
    project = seed_demo_project(session)

    client, calls = make_mock_client({
        LOG4J_PURL: {"vulns": [GHSA_LOG4J]},
        FASTJSON_PURL: {"vulns": [dict(CVE_FASTJSON)]},
    })
    result = run_full_pipeline(
        session, project.id, out_dir=tmp_path / "out2", osv_client=client,
    )

    assert len(calls) >= 2
    vuln_reqs = [r for r in result.requirements if r.category == "第三方组件风险"]
    assert {r.template_id for r in vuln_reqs} == {"SEC-V14-801"}

    cve_rows = [
        v for v in session.query(VulnerabilityRecord).all()
        if v.cve_id.startswith("CVE-")
    ]
    assert {v.cve_id for v in cve_rows} >= {"CVE-2021-44228"}

    text = "\n".join(iter_document_text(Document(str(result.documents["sbom_vuln"]))))
    assert "CVE-2021-44228" in text
    session.close()
