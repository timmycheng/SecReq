# -*- coding: utf-8 -*-
"""《安全需求说明书》导出(#162): 章节结构、执行摘要与定级来源标注。

用 python-docx 读回生成的字节流, 校验章节文本与关键内容。
"""
from conftest import add_base_project
from models import Filing, GradingSurvey, SecurityRequirement, System
from services.doc_export import build_full_docx

import io

from docx import Document


def _make_project(session, with_filing=True):
    project = add_base_project(session)
    if with_filing:
        filing = Filing(name="网银核心备案", code="BA-T1", level="三级")
        session.add(filing)
        session.flush()
        system = System(name="手机银行系统", code="SYS-T1", filing_id=filing.id)
        session.add(system)
        session.flush()
        project.system_id = system.id
    session.add(GradingSurvey(
        project_id=project.id, answers_json=[], suggested_level="二级",
        final_level="三级", manual_adjust_note="涉及交易类敏感数据"))
    session.flush()
    return project


def _add_req(session, project, req_id="SEC-XXX-001", priority="critical",
             title="关键需求", category="数据安全", source="数据资产:客户信息表"):
    session.add(SecurityRequirement(
        project_id=project.id, req_id=req_id, template_id="SEC-XXX-001",
        title=title, description="需对客户信息表实施脱敏存储", category=category,
        priority=priority, acceptance_criteria="抽查字段全部脱敏",
        suggested_phase="design", source_entity_type="data_asset", source_entity_id=1,
        source_entity_uid="uid-1", source_label=source,
        trigger_reason="命中 3级_C2 敏感资产",
        regulatory_ref=[{"file": "个人金融信息保护技术规范", "clause": "第6.1条",
                         "summary": "脱敏要求"}],
    ))
    session.commit()


def _doc_text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class _FakeVuln:
    severity = "low"
    cve_id = "CVE-2026-0001"
    cnnvd_id = None
    component_id = 999
    affected_range = "<1.0"
    fix_version = "1.0"
    summary = "测试"
    source = "osv_local"


def test_docx_structure_and_executive_summary(session):
    """目录/执行摘要/类目分章存在; 结论按 critical 分档为红色档。"""
    project = _make_project(session)
    _add_req(session, project)
    reqs = session.query(SecurityRequirement).filter_by(project_id=project.id).all()

    content = build_full_docx(project, reqs, [_FakeVuln()], components=[])
    text = _doc_text(content)

    for chapter in ["目录", "一、执行摘要", "二、项目概况与定级",
                    "三、安全需求清单", "3.1 数据安全", "四、漏洞清单"]:
        assert chapter in text, f"缺少章节: {chapter}"
    assert "不建议直接通过" in text          # critical 存在 → 红档结论
    assert "Top 风险" in text and "关键需求" in text
    assert "未发现漏洞记录" not in text       # 传入了漏洞


def test_docx_grading_source_from_filing(session):
    """挂备案的系统: 概况表展示所属系统与备案定级来源。"""
    project = _make_project(session, with_filing=True)
    content = build_full_docx(project, [], [], components=[])
    text = _doc_text(content)
    assert "所属系统" in text and "手机银行系统" in text
    assert "定级备案《网银核心备案》等保三级" in text
    assert "未归属" not in text


def test_docx_diff_chapter_numbering(session):
    """无上一轮: 章节到四为止; 有 diff_data: 五、与上一轮差异出现。"""
    project = _make_project(session, with_filing=False)
    _add_req(session, project, priority="high")
    reqs = session.query(SecurityRequirement).filter_by(project_id=project.id).all()

    text = _doc_text(build_full_docx(project, reqs, [], components=[]))
    assert "五、与上一轮差异" not in text

    diff = {
        "previous_project": {"project_id": 0, "project_name": "P", "project_code": "XM-1",
                             "created_at": None},
        "added": [], "removed": [],
        "changed": [{"fields": ["priority"],
                     "previous": {"req_id": "SEC-XXX-001", "title": "关键需求",
                                  "priority": "critical", "category": "数据安全",
                                  "source_label": None, "status": "open",
                                  "suggested_phase": "design"},
                     "current": {"req_id": "SEC-XXX-001", "title": "关键需求",
                                 "priority": "high", "category": "数据安全",
                                 "source_label": None, "status": "open",
                                 "suggested_phase": "design"}}],
        "summary": {"current_total": 1, "previous_total": 1, "added": 0,
                    "removed": 0, "changed": 1},
    }
    text = _doc_text(build_full_docx(project, reqs, [], components=[], diff_data=diff))
    assert "五、与上一轮(XM-1)差异" in text
    assert "变更" in text and "priority" in text
