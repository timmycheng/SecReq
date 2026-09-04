# -*- coding: utf-8 -*-
"""「下载 Word 文档」产物导出(走查整改: 全文下载用 .docx, 分区粘贴用前端 HTML 剪贴板)。

生成《安全需求说明书》, 结构对齐"结论先行"的审阅动线(#162):
封面/目录 → 一、执行摘要(自动结论+关键数字+Top风险+合规覆盖)
→ 二、项目概况与定级(含所属系统与定级来源)
→ 三、安全需求清单(按类目分章, 每章小结+表格)
→ 四、漏洞清单 → 五、与上一轮差异(评估继承场景, 可选)。
中文字体: 标题黑体、正文宋体; 表格带边框, 表头灰底。
"""
import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import shared.constants as C

_HEADER_FILL = "EDEDED"
_CRITICAL_RED = RGBColor(0xC0, 0x00, 0x00)
_ORANGE = RGBColor(0xD4, 0x69, 0x00)
_GREEN = RGBColor(0x1E, 0x7D, 0x32)
_GREY = RGBColor(0x59, 0x59, 0x59)

_PRIORITY_ORDER = {p: i for i, p in enumerate(["critical", "high", "medium", "low"])}


def _set_cn_font(run, name: str = "宋体", size: float = 10.5, bold: bool = False,
                 color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), name)


def _shade(cell, fill: str) -> None:
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_text(cell, text: str, bold: bool = False, size: float = 9, red: bool = False) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text or "—")
    _set_cn_font(run, size=size, bold=bold, color=_CRITICAL_RED if red else None)


def _add_table(doc: Document, headers: list[str], widths_cm: list[float]) -> list:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (header, width) in enumerate(zip(headers, widths_cm)):
        cell = table.rows[0].cells[i]
        _cell_text(cell, header, bold=True)
        _shade(cell, _HEADER_FILL)
        cell.width = Cm(width)
    return table


def _heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_cn_font(run, name="黑体", size=14, bold=True)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(6)


def _para(doc: Document, text: str, size: float = 10.5, bold: bool = False,
          color: RGBColor | None = None) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_cn_font(run, size=size, bold=bold, color=color)
    para.paragraph_format.space_after = Pt(4)


def _note(doc: Document, text: str) -> None:
    """小号灰色说明段(数据来源声明、覆盖缺口提示)。"""
    _para(doc, text, size=8.5, color=_GREY)


#: 数据来源 → 人读文案(导出文档里必须交代来源, 便于合规说明)
_SOURCE_LABELS = {
    "osv_local": "本地离线漏洞库",
    "osv_online": "OSV.dev 在线库",
    "sca": "行内 SCA 平台",
}


def _vuln_source_note(doc: Document, vulnerabilities: list) -> None:
    """标注数据来源与库版本 —— 内网交付时这是合规说明的一部分。"""
    sources = sorted({getattr(v, "source", "osv_local") or "osv_local" for v in vulnerabilities})
    text = "数据来源: " + "、".join(_SOURCE_LABELS.get(s, s) for s in sources)
    try:
        from services.vulndb import VulnDb
        meta = VulnDb().meta()
        text += f" v{meta.get('db_version', '未知版本')}(构建于 {meta.get('built_at', '未知')})"
    except Exception:  # 库不可读时只显示来源, 不阻断导出
        pass
    _note(doc, text)
    if any(getattr(v, "source", "") == "osv_local" for v in vulnerabilities):
        _note(doc, "本结果由本地离线漏洞库匹配得出, 未访问任何外部服务。")


def _uncovered_note(doc: Document, components: list) -> None:
    """把"未覆盖 / 无法判定 / 命中但待确认"的组件单独列出。

    三者都不能混进"未发现漏洞"里 —— 那会给人虚假的安全感。
    """
    pending = [
        c for c in components
        if getattr(c, "vuln_status", None) in ("not_covered", "undetermined")
        # 命中但带说明: 跨渠道模糊匹配、麒麟推断、版本缺修订号等
        or (getattr(c, "vuln_status", None) == "hit" and getattr(c, "vuln_status_note", None))
    ]
    if not pending:
        return
    lines = []
    for comp in pending:
        label = C.VULN_QUERY_STATUS.get(comp.vuln_status, comp.vuln_status)
        if comp.vuln_status == "hit":
            label = "命中待确认"
        lines.append(f"{comp.name}@{comp.version}: {label}。{comp.vuln_status_note or ''}")
    _note(doc, "以下组件需人工确认: " + " ".join(lines))


def _conclusion(requirements: list, vulnerabilities: list) -> tuple[str, str, RGBColor]:
    """自动结论(与结果页执行摘要同口径): 按 critical/high 的需求与漏洞分档。"""
    crit = sum(1 for r in requirements if r.priority == "critical")
    high = sum(1 for r in requirements if r.priority == "high")
    crit_v = sum(1 for v in vulnerabilities if v.severity == "critical")
    high_v = sum(1 for v in vulnerabilities if v.severity == "high")
    if crit or crit_v:
        return (f"不建议直接通过: 存在 {crit} 条关键需求与 {crit_v} 个严重漏洞",
                "关键项为硬性安全要求, 建议整改闭环后复评; 优先处理下表 Top 风险。",
                _CRITICAL_RED)
    if high or high_v:
        return (f"有条件通过: 无关键(critical)项, 有 {high} 条高优先级需求与 {high_v} 个高危漏洞",
                "建议按 Top 风险排期整改, 其余需求按建议阶段落实。",
                _ORANGE)
    return (f"基线整体可控: 共 {len(requirements)} 条需求, 均非 critical/high",
            "按建议阶段落实即可, 无需额外整改决策。",
            _GREEN)


def _ref_files(requirement) -> list[str]:
    return [ref.get("file", "") for ref in (getattr(requirement, "regulatory_ref", None) or [])
            if ref.get("file")]


def _executive_summary(doc: Document, project, requirements: list,
                       vulnerabilities: list) -> None:
    """一、执行摘要: 结论先行, 30 秒回答"能不能过、先看什么"(#162)。"""
    _heading(doc, "一、执行摘要")
    text, detail, color = _conclusion(requirements, vulnerabilities)
    _para(doc, text, size=12, bold=True, color=color)
    _para(doc, detail, size=9.5, color=_GREY)

    crit = sum(1 for r in requirements if r.priority == "critical")
    high = sum(1 for r in requirements if r.priority == "high")
    confirmed = sum(1 for r in requirements if getattr(r, "reg_confirmed", False))
    crit_v = sum(1 for v in vulnerabilities if v.severity == "critical")
    high_v = sum(1 for v in vulnerabilities if v.severity == "high")
    numbers = _add_table(doc,
                         ["安全需求", "紧急(critical)", "高(high)", "已确认", "严重漏洞", "高危漏洞"],
                         [2.4] * 6)
    values = [str(len(requirements)), str(crit), str(high), str(confirmed),
              str(crit_v), str(high_v)]
    row = numbers.add_row().cells
    for i, value in enumerate(values):
        _cell_text(row[i], value, bold=(i == 0), red=(i in (1, 4) and value != "0"))

    top = sorted(
        (r for r in requirements if r.priority in ("critical", "high")),
        key=lambda r: (_PRIORITY_ORDER.get(r.priority, 9), r.req_id),
    )[:5]
    if top:
        _para(doc, "Top 风险(完整清单见第三章):", size=10.5, bold=True)
        table = _add_table(doc, ["编号", "优先级", "需求标题", "来源"], [3.4, 1.8, 7.4, 4.8])
        for r in top:
            row = table.add_row().cells
            _cell_text(row[0], r.req_id, red=r.priority == "critical")
            _cell_text(row[1], C.label(C.REQUIREMENT_PRIORITY_LABELS, r.priority),
                       red=r.priority == "critical")
            _cell_text(row[2], r.title)
            _cell_text(row[3], getattr(r, "source_label", None) or "—")

    if project.compliance_targets:
        lines = []
        for code in project.compliance_targets:
            keyword = C.COMPLIANCE_FILE_KEYWORDS.get(code)
            label = C.label(C.COMPLIANCE_TARGETS, code)
            count = sum(1 for r in requirements if keyword
                        and any(keyword in f for f in _ref_files(r)))
            lines.append(f"{label}: {count} 条" if count else f"{label}: 未直接命中")
        _para(doc, "合规目标覆盖: " + ";".join(lines) + "(按需求监管出处统计)。", size=9.5)


def _requirements_by_category(doc: Document, requirements: list) -> None:
    """三、安全需求清单: 按类目分章, 每章一句小结 + 减负表格(类目列省去)。"""
    _heading(doc, f"三、安全需求清单(共 {len(requirements)} 条, 按类目分章)")
    if not requirements:
        doc.add_paragraph("尚未生成安全需求, 请先在向导确认页执行「生成安全基线」。")
        return

    groups: dict[str, list] = {}
    for r in sorted(requirements, key=lambda r: (_PRIORITY_ORDER.get(r.priority, 9), r.req_id)):
        groups.setdefault(r.category, []).append(r)
    label_values = list(C.TRIGGER_CATEGORY_LABELS.values())
    ordered = sorted(groups.items(),
                     key=lambda kv: label_values.index(kv[0]) if kv[0] in label_values else 99)

    for idx, (label, rows) in enumerate(ordered, start=1):
        crit = sum(1 for r in rows if r.priority == "critical")
        high = sum(1 for r in rows if r.priority == "high")
        _heading(doc, f"3.{idx} {label}(共 {len(rows)} 条)")
        summary = f"其中紧急(critical){crit} 条、高(high){high} 条。" if crit or high \
            else "无紧急/高优先级条目。"
        sources = {getattr(r, "source_label", None) or "" for r in rows} - {""}
        if sources:
            summary += f"主要来源: {'、'.join(sorted(sources)[:3])}。"
        _para(doc, summary, size=9.5, color=_GREY)

        table = _add_table(doc,
                           ["序号", "需求内容", "优先级", "来源", "验收标准", "合规依据/确认"],
                           [1.0, 7.0, 1.5, 2.6, 3.4, 2.9])
        for no, req in enumerate(rows, start=1):
            row = table.add_row().cells
            _cell_text(row[0], str(no), red=req.priority == "critical")

            content = row[1]
            content.text = ""
            p1 = content.paragraphs[0]
            _set_cn_font(p1.add_run(f"[{req.req_id}] {req.title}"), bold=True)
            p2 = content.add_paragraph()
            _set_cn_font(p2.add_run(req.description or ""))

            _cell_text(row[2], C.label(C.REQUIREMENT_PRIORITY_LABELS, req.priority),
                       red=req.priority == "critical")
            _cell_text(row[3], getattr(req, "source_label", None)
                       or (f"{req.source_entity_type}#{req.source_entity_id}"
                           if getattr(req, "source_entity_type", None) else "—"))
            _cell_text(row[4], req.acceptance_criteria or "—")
            refs = ";\n".join(f"《{ref.get('file', '')}》{ref.get('clause', '')}"
                              for ref in (getattr(req, "regulatory_ref", None) or [])
                              if ref.get("file"))
            confirmed = "✓ 已确认" if getattr(req, "reg_confirmed", False) else "□ 未确认"
            _cell_text(row[5], (refs or "—") + f"\n{confirmed}", size=8.5)


def build_full_docx(
    project,
    requirements: list,
    vulnerabilities: list,
    components: list | None = None,
    diff_data: dict | None = None,
) -> bytes:
    """生成整册《安全需求说明书》.docx 字节流。

    diff_data: services.requirement_diff.diff_requirements 的结果(评估继承场景,
    有上一轮时附「与上一轮差异」章节); None 则不出现该章节。
    """
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)  # A4 纵向

    # 封面标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cn_font(title.add_run(f"{project.name} 安全需求说明书"), name="黑体", size=20, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cn_font(
        sub.add_run(f"项目编码: {project.code}    导出时间: {datetime.now():%Y-%m-%d %H:%M}"),
        size=9,
    )

    # 目录(静态章节清单)
    _heading(doc, "目录")
    for entry in ["一、执行摘要", "二、项目概况与定级", "三、安全需求清单(按类目分章)",
                  "四、漏洞清单", *(["五、与上一轮差异"] if diff_data is not None else [])]:
        _para(doc, entry, size=10.5)

    # 一、执行摘要
    _executive_summary(doc, project, requirements, vulnerabilities)

    # 二、项目概况与定级
    _heading(doc, "二、项目概况与定级")
    survey = project_survey(project)
    survey_obj = getattr(project, "survey", None)
    level = survey_obj.effective_level() if survey_obj else ""
    system = getattr(project, "system", None)
    filing = getattr(system, "filing", None) if system else None
    grading_source = (f"定级备案《{filing.name}》等保{filing.level}(系统挂靠继承)"
                      if filing else (survey or "—"))
    meta_rows = [
        ("项目名称", project.name),
        ("项目编码", project.code),
        ("所属系统", system.name if system else "未归属(可在项目中编辑归属)"),
        ("有效定级", f"等保{level}" if level else "未定级"),
        ("定级来源", grading_source),
        ("项目类型", "、".join(
            C.label(C.PROJECT_TYPES, t) for t in (getattr(project, "types", None) or [])) or "—"),
        ("合规目标", "、".join(
            C.label(C.COMPLIANCE_TARGETS, t) for t in (project.compliance_targets or [])) or "—"),
        ("定级结论", survey or "—"),
        ("安全需求数", f"{len(requirements)} 条"),
    ]
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.style = "Table Grid"
    for key, value in meta_rows:
        row = meta_table.add_row()
        _cell_text(row.cells[0], key, bold=True)
        _shade(row.cells[0], _HEADER_FILL)
        _cell_text(row.cells[1], value, size=10)
        row.cells[0].width, row.cells[1].width = Cm(3.5), Cm(13.0)

    # 三、安全需求清单(按类目分章)
    _requirements_by_category(doc, requirements)

    # 四、漏洞清单
    _heading(doc, f"四、漏洞清单(共 {len(vulnerabilities)} 条)")
    if vulnerabilities:
        comp_by_id = {c.id: c for c in (components or [])}
        table = _add_table(
            doc,
            # v2.2.0: 补 CNNVD 编号 —— 银行合规通报常要求国产编号, 事后手工补录成本很高
            ["等级", "CVE", "CNNVD", "组件", "受影响范围", "修复版本", "简述"],
            [1.3, 2.8, 2.6, 3.3, 3.1, 2.4, 3.6],
        )
        for v in vulnerabilities:
            comp = comp_by_id.get(v.component_id)
            comp_label = f"{comp.name}@{comp.version}" if comp else f"组件#{v.component_id}"
            row = table.add_row().cells
            high = v.severity in ("critical", "high")
            _cell_text(row[0], C.label(C.SEVERITY_LABELS, v.severity), red=high)
            _cell_text(row[1], v.cve_id, red=high)
            _cell_text(row[2], getattr(v, "cnnvd_id", None) or "—")
            _cell_text(row[3], comp_label)
            _cell_text(row[4], v.affected_range or "—")
            _cell_text(row[5], v.fix_version or "官方暂未发布修复版")
            _cell_text(row[6], v.summary or "")
        _vuln_source_note(doc, vulnerabilities)
        _uncovered_note(doc, components or [])
    else:
        doc.add_paragraph("未发现漏洞记录。")

    # 五、与上一轮差异(评估继承场景)
    if diff_data is not None:
        _diff_chapter(doc, diff_data)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _diff_chapter(doc: Document, diff: dict) -> None:
    """「与上一轮差异」章节(评估继承 #151): 分期建设场景下审阅者最关心增量。"""
    prev_code = (diff.get("previous_project") or {}).get("project_code") or "上一轮"
    added_rows = diff.get("added") or []
    removed_rows = diff.get("removed") or []
    changed_rows = diff.get("changed") or []
    _heading(doc, f"五、与上一轮({prev_code})差异")
    if not (added_rows or removed_rows or changed_rows):
        doc.add_paragraph(f"与上一轮({prev_code})对比, 安全需求无变化。")
        return
    para = doc.add_paragraph()
    _set_cn_font(para.add_run(
        f"新增 {len(added_rows)} 条;移除 {len(removed_rows)} 条;变更 {len(changed_rows)} 条。"),
        bold=True,
    )
    table = _add_table(doc, ["变更", "编号", "优先级", "需求标题", "来源", "说明"],
                       [1.4, 3.0, 1.4, 5.4, 3.0, 3.3])
    for kind, rows, color in (("新增", added_rows, None), ("移除", removed_rows, _CRITICAL_RED)):
        for r in rows:
            row = table.add_row().cells
            _cell_text(row[0], kind, bold=True, red=color is not None)
            _cell_text(row[1], r["req_id"])
            _cell_text(row[2], C.label(C.REQUIREMENT_PRIORITY_LABELS, r["priority"]))
            _cell_text(row[3], r["title"])
            _cell_text(row[4], r.get("source_label") or "—")
            _cell_text(row[5], "本轮基线不再包含" if kind == "移除" else "本轮基线新增要求")
    for c in changed_rows:
        row = table.add_row().cells
        cur = c.get("current") or {}
        _cell_text(row[0], "变更", bold=True)
        _cell_text(row[1], cur.get("req_id", ""))
        _cell_text(row[2], C.label(C.REQUIREMENT_PRIORITY_LABELS, cur.get("priority", "")))
        _cell_text(row[3], cur.get("title", ""))
        _cell_text(row[4], cur.get("source_label") or "—")
        _cell_text(row[5], "变化字段: " + "、".join(c.get("fields") or []))
    _note(doc, "差异按知识库模板与来源实体对齐(template_id + source_entity_uid); "
               "同一输入实体的要求调整记为「变更」。")


def project_survey(project) -> str | None:
    """定级结论一句话(人工修正标注); 无问卷返回 None。"""
    survey = getattr(project, "survey", None)
    if survey is None or not survey.effective_level():
        return None
    if survey.final_level and survey.suggested_level:
        tag = "人工修正"
    elif survey.final_level:
        tag = "直接指定"
    else:
        tag = "系统建议"
    return f"等保{survey.effective_level()}({tag})"
