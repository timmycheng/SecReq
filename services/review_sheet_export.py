# -*- coding: utf-8 -*-
"""《项目安全评审表》导出(#230): 评审会可直接归档上会的第 5 份文档。

内容: 门禁状态、需求覆盖统计、漏洞概况、遗留问题、评审意见与签字栏。
数据来源与评审中心(/review/state)一致; 仅门禁推进到终审环节的项目可导出。
"""
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from models import ReviewGate
from services.doc_export import _heading, _para, _set_cn_font

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

_GATE_STATUS_LABELS = {
    "pending": "待提交", "in_review": "评审中", "passed": "终审通过",
    "rejected": "已否决", "rectifying": "退回整改中",
}
_LIFECYCLE_LABELS = {
    "open": "待确认", "confirmed": "已确认",
    "reviewed": "评审通过", "rectifying": "整改中",
}


def build_review_sheet_docx(project, gate: ReviewGate | None,
                            requirement_summary: dict,
                            evidences: list[dict],
                            chain_valid: bool,
                            vuln_summary: dict | None = None) -> bytes:
    """生成《项目安全评审表》.docx 字节流。

    gate/evidences 来自 services.review_service.review_state;
    vuln_summary: {"total": n, "critical": n, "high": n}, 可为 None(未查询)。
    """
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cn_font(title.add_run(f"{project.name} 项目安全评审表"), name="黑体", size=20, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cn_font(sub.add_run(
        f"项目编码: {project.code}    导出时间: {datetime.now():%Y-%m-%d %H:%M}"), size=9)

    # 一、门禁状态
    _heading(doc, "一、评审门禁状态")
    if gate is None:
        _para(doc, "尚未提交评审。", size=10.5)
    else:
        _para(doc, "门禁类型: 需求门禁(requirement)", size=10.5)
        _para(doc, f"当前状态: {_GATE_STATUS_LABELS.get(gate['status'], gate['status'])}"
                   f"({gate.get('status_verb') or ''})", size=10.5)
        _para(doc, f"提交时间: {gate.get('submitted_at') or '—'}    "
                   f"评审时间: {gate.get('reviewed_at') or '—'}    "
                   f"终审时间: {gate.get('final_reviewed_at') or '—'}", size=10.5)
        conclusion = gate.get("reviewer_conclusion")
        _para(doc, "评审员裁定: " + {
            "approve": "通过", "reject": "否决", "request_change": "退回整改",
        }.get(conclusion, "—"), size=10.5)
        if gate.get("reviewer_opinion"):
            _para(doc, f"评审意见: {gate['reviewer_opinion']}", size=10.5)
        if gate.get("final_opinion"):
            _para(doc, f"终审意见: {gate['final_opinion']}", size=10.5)
        _para(doc, "评审留痕链校验: " + ("完整" if chain_valid else "发现篡改, 请立即核查!"),
              size=10.5)

    # 二、需求覆盖统计
    _heading(doc, "二、安全需求覆盖统计")
    _para(doc, f"待确认: {requirement_summary.get('open', 0)} 条    "
               f"已确认: {requirement_summary.get('confirmed', 0)} 条    "
               f"评审通过: {requirement_summary.get('reviewed', 0)} 条    "
               f"整改中: {requirement_summary.get('rectifying', 0)} 条", size=10.5)
    total = sum(requirement_summary.values())
    _para(doc, f"需求总数: {total} 条, 其中评审通过率 "
               f"{(requirement_summary.get('reviewed', 0) / total * 100 if total else 0):.0f}%",
          size=10.5)

    # 三、漏洞概况
    _heading(doc, "三、第三方组件漏洞概况")
    if vuln_summary is None:
        _para(doc, "未执行漏洞查询(离线模式或未维护组件清单)。", size=10.5)
    else:
        _para(doc, f"漏洞记录 {vuln_summary.get('total', 0)} 条, 其中严重 "
                   f"{vuln_summary.get('critical', 0)} 条、高危 "
                   f"{vuln_summary.get('high', 0)} 条。", size=10.5)

    # 四、遗留问题
    _heading(doc, "四、遗留问题与整改事项")
    rectifying = requirement_summary.get("rectifying", 0)
    open_count = requirement_summary.get("open", 0)
    if rectifying == 0 and open_count == 0 and gate and gate["status"] == "passed":
        _para(doc, "无。全部需求已评审通过, 本轮基线已具备写回条件。", size=10.5)
    else:
        if open_count:
            _para(doc, f"- {open_count} 条需求待项目经理确认。", size=10.5)
        if rectifying:
            _para(doc, f"- {rectifying} 条需求处于整改中, 需整改后重新提交复审。", size=10.5)
        if gate and gate["status"] == "rectifying":
            _para(doc, "- 项目门禁整体处于退回整改状态, 整改完成后重新提交评审。", size=10.5)
        if gate and gate["status"] == "rejected":
            _para(doc, "- 项目门禁已被否决, 需按评审意见整改后重新提交。", size=10.5)

    # 五、评审动作留痕
    _heading(doc, "五、评审动作留痕")
    if not evidences:
        _para(doc, "暂无评审动作记录。", size=10.5)
    for ev in evidences:
        line = (f"{ev['timestamp'][:19].replace('T', ' ')}  "
                f"{ev['action']}  "
                + (f"意见: {ev['comment']}" if ev.get("comment") else ""))
        _para(doc, line, size=9.5)

    # 六、评审意见与签字栏
    _heading(doc, "六、评审意见与签字栏")
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    heads = [("项目经理", "签字/日期"), ("安全评审员", "签字/日期"),
             ("安全负责人", "签字/日期"), ("备注", "评审结论(通过/有条件通过/不通过)")]
    for i, (left, right) in enumerate(heads):
        row = table.rows[i]
        _set_cn_font(row.cells[0].paragraphs[0].add_run(left), size=10)
        _set_cn_font(row.cells[2].paragraphs[0].add_run(right), size=10)
    for row in table.rows:
        row.height = Cm(1.2)

    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
