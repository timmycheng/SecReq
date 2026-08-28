# -*- coding: utf-8 -*-
"""Word 文档生成服务(python-docx)。

DESIGN.md 模块4(升级版): 按行内模板生成 5 份中文 Word 文档, 均含封面
(项目名/编码/生成时间/编制人/审核人签字栏):

1. grading     《系统定级建议书》: 问卷答案表→定级结论→判定理由→判定依据→安全中心复核意见
2. requirement 《需求规格说明书-安全需求章节》: 按"监管报送类/等保条款类/通用安全类"分组的
               需求表格(含合规依据列) + 权限矩阵附录 + 评审记录页
3. design      《总体设计说明书-安全设计章节》: 软件版本清单/数据字典(JR/T 五级)/
               API属性/资产清单/密码策略/认证方式/监管报送事项清单
4. sbom_vuln   《SBOM及漏洞清单》: SBOM组件表 + 漏洞表(高危标红置顶)
5. review      《项目安全评审表》: 各门禁状态/需求覆盖统计/漏洞概况/遗留问题(评审会材料)

版式参数(字体字号/标红颜色/表头底纹)从 docs/templates/doc_style.yml 读取,
由安全中心维护替换, 修改不涉及代码变更。
"""
import logging
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor

import shared.constants as C
from models import SecurityRequirement
from rules.context import RequirementContext
from rules.policy import effective_password_policy

logger = logging.getLogger(__name__)

DEFAULT_STYLE = {
    "fonts": {
        "body_cn": "宋体",
        "heading_cn": "黑体",
        "body_size_pt": 10.5,
        "table_size_pt": 9,
        "heading1_size_pt": 16,
        "heading2_size_pt": 13,
    },
    "cover": {"title_size_pt": 26},
    "colors": {"vuln_red": "C00000", "header_shading": "D9E2F3"},
    "signature_roles": ["编制人", "审核人"],
}

STYLE_PATH = Path(__file__).resolve().parent.parent / "docs" / "templates" / "doc_style.yml"

# OWASP ASVS 4.0.3 章节 → 中文名(需求文档分组用)
ASVS_CHAPTERS = {
    "V1": "架构、设计与威胁建模",
    "V2": "身份验证",
    "V3": "会话管理",
    "V4": "访问控制",
    "V5": "验证、净化与输出编码",
    "V6": "加密的存储",
    "V7": "错误处理与日志记录",
    "V8": "数据保护",
    "V9": "通信",
    "V10": "恶意代码",
    "V11": "反自动化与业务逻辑",
    "V12": "文件与资源",
    "V13": "API 与 Web 服务",
    "V14": "配置",
}

# 认证方式 → 设计说明要点(总体设计说明书第五节)
AUTH_METHOD_DESIGN_NOTES = {
    "password": "采用用户名口令登录; 口令强度、有效期、错误锁定等参数按本项目密码策略基线执行, 口令仅以不可逆摘要存储。",
    "sms_otp": "短信验证码设计为6位数字、有效期60秒、单次有效; 同一号码限频防轰炸(如60秒内1条/每日上限), 校验失败次数受限并记录审计日志。",
    "dynamic_otp": "接入OTP动态口令(TOTP), 密钥开通时通过安全渠道扫码绑定, 并提示用户保存恢复码。",
    "third_oauth": "第三方登录采用授权码模式; 回调地址严格白名单校验, state 参数防CSRF, access_token/openid 在服务端校验真伪后方可建会话。",
    "sso": "对接行内统一认证(SSO); 会话票据签发/注销与统一认证中心联动, 行内单点登出需同步销毁本地会话。",
    "biometric": "生物识别仅作为辅助认证因子, 原始生物特征模板加密存储且不落业务日志, 提供替代验证方式。",
}


def load_doc_style(path: str | Path | None = None) -> dict:
    """合并默认值加载版式配置; 配置文件可整体缺失或部分覆盖。"""
    style = {
        section: (dict(values) if isinstance(values, dict) else values)
        for section, values in DEFAULT_STYLE.items()
    }
    path = Path(path or STYLE_PATH)
    if path.exists():
        with open(path, encoding="utf-8") as f:
            override = yaml.safe_load(f) or {}
        for section, values in override.items():
            if isinstance(values, dict) and isinstance(style.get(section), dict):
                style[section].update(values)
            else:
                style[section] = values
    return style


# ────────────────────────── 底层排版辅助 ──────────────────────────


class DocBuilder:
    """对 python-docx 的中文文档封装: 统一字体/标题/表格/封面页脚。"""

    def __init__(self, style: dict):
        self.style = style
        self.doc = Document()
        f = style["fonts"]
        self._setup_styles(f)

    # ── 样式初始化 ──

    def _setup_styles(self, fonts: dict) -> None:
        normal = self.doc.styles["Normal"]
        normal.font.name = fonts["body_cn"]
        normal.font.size = Pt(fonts["body_size_pt"])
        _style_east_asia(normal, fonts["body_cn"])

        for level, key in ((1, "heading1_size_pt"), (2, "heading2_size_pt")):
            st = self.doc.styles[f"Heading {level}"]
            st.font.name = fonts["heading_cn"]
            st.font.size = Pt(fonts[key])
            st.font.color.rgb = RGBColor(0, 0, 0)
            _style_east_asia(st, fonts["heading_cn"])

    # ── 结构元素 ──

    def cover(
        self,
        title: str,
        project_name: str,
        project_code: str,
        generated_at_text: str,
        preparer: str = "",
    ) -> None:
        """封面: 居中大标题 + 元信息 + 签字栏(编制人/审核人), 之后分页。"""
        fonts = self.style["fonts"]
        title_size = self.style["cover"]["title_size_pt"]
        roles = self.style["signature_roles"] or DEFAULT_STYLE["signature_roles"]

        for _ in range(6):  # 视觉上下移
            self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, title, size=title_size, bold=True, font=fonts["heading_cn"])

        self.doc.add_paragraph()
        meta_lines = [
            f"项目名称：{project_name}",
            f"项目编码：{project_code}",
            f"生成时间：{generated_at_text}",
        ]
        for line in meta_lines:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(para, line, size=14)

        for _ in range(4):
            self.doc.add_paragraph()

        # 签字栏: 编制人姓名已填, 审核人留待手工签字
        sign_table = self.doc.add_table(rows=1, cols=4)
        pairs = [
            (f"{roles[0]}：", preparer or "（待填）"),
            (f"{roles[1]}：", "_" * 12),
        ]
        cells = sign_table.rows[0].cells
        for idx, (label, value) in enumerate(pairs):
            cells[idx * 2].text = ""
            cells[idx * 2 + 1].text = ""
            _run(cells[idx * 2].paragraphs[0], label, size=12)
            _run(cells[idx * 2 + 1].paragraphs[0], value, size=12)

        self.doc.add_page_break()

    def heading(self, text: str, level: int = 1) -> None:
        self.doc.add_heading(text, level=level)

    def para(self, text: str, bold: bool = False, align=None, size: float | None = None) -> None:
        p = self.doc.add_paragraph()
        if align is not None:
            p.alignment = align
        _run(p, text, bold=bold, size=size)

    def bullets(self, items: list[str]) -> None:
        for item in items:
            p = self.doc.add_paragraph(style="List Bullet")
            _run(p, item)

    def table(
        self,
        headers: list[str],
        rows: list[list],
        red_cells: set[tuple[int, int]] | None = None,
    ) -> None:
        """带边框表格; rows 单元格可为字符串或 [行文本...] 多段; red_cells 红色字体(r,c)。"""
        red_cells = red_cells or set()
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        _shade_header(table.rows[0], headers, self.style)

        for r_idx, row_values in enumerate(rows):
            row = table.add_row()
            for c_idx, value in enumerate(row_values):
                cell = row.cells[c_idx]
                cell.text = ""
                lines = value if isinstance(value, (list, tuple)) else [value]
                for line_no, line in enumerate(lines):
                    para = cell.paragraphs[0] if line_no == 0 else cell.add_paragraph()
                    color = self.style["colors"]["vuln_red"] if (r_idx, c_idx) in red_cells else None
                    _run(para, str(line), color=color, size=self.style["fonts"]["table_size_pt"])
        self.doc.add_paragraph()  # 表后留白

    def save(self, path: str | Path) -> Path:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        _footer_page_number(self.doc)
        self.doc.save(path)
        return path


def _run(paragraph, text: str, *, bold=False, size=None, font=None, color: str | None = None):
    run = paragraph.add_run(text)
    fonts = DEFAULT_STYLE["fonts"]
    run.font.name = font or fonts["body_cn"]
    run.font.size = Pt(size or fonts["body_size_pt"])
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font or fonts["body_cn"])
    return run


def _style_east_asia(style_obj, cn_font: str) -> None:
    """样式级中文字体必须写 eastAsia 属性, 否则 Word 中文回落到默认西文字体。"""
    rpr = style_obj.element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), cn_font)


def _shade_header(row, headers: list[str], style: dict) -> None:
    fill = style["colors"]["header_shading"]
    size = style["fonts"]["table_size_pt"]
    for cell, text in zip(row.cells, headers):
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, text, bold=True, size=size)


def _footer_page_number(doc: Document) -> None:
    footer_p = doc.sections[-1].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head = footer_p.add_run("第 ")
    head.font.size = Pt(9)
    tail = footer_p.add_run(" 页")
    tail.font.size = Pt(9)
    # 把 PAGE 域插到两段文字中间
    page_run = _field(footer_p, "PAGE")
    tail._element.addprevious(page_run)


def _field(paragraph, instr: str):
    """构造 Word 域运行元素(页码 PAGE), 返回待挂载的 <w:r>。"""
    run = paragraph.add_run()
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    element = OxmlElement("w:instrText")
    element.set(qn("xml:space"), "preserve")
    element.text = f" {instr} "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._element.extend([begin, element, end])
    return run._element


def _fmt_date(dt) -> str:
    return dt.strftime("%Y年%m月%d日 %H:%M")


def _yesno(flag: bool) -> str:
    return "是" if flag else "否"


# ────────────────────────── 文档一: 系统定级建议书 ──────────────────────────


def build_grading_report(ctx: RequirementContext, builder: DocBuilder,
                         gates: list | None = None) -> None:
    project = ctx.project
    survey = ctx.survey

    builder.heading("一、项目基本信息")
    builder.table(
        ["项目", "内容"],
        [
            ["项目名称", project.name],
            ["项目编码", project.code],
            ["项目类型", C.label(C.PROJECT_TYPES, project.type)],
            ["所属业务条目", project.industry],
            ["用户规模", ctx.user_scale_text],
            ["部署环境", "、".join(C.label(C.DEPLOY_ENVS, e) for e in (project.deploy_env or []))],
            ["是否涉及公网访问", _yesno(project.is_public)],
            ["项目经理", project.pm_name or ""],
            ["开发负责人", project.dev_lead_name or ""],
            ["安全对接人", project.sec_contact_name or ""],
            ["合规目标", "、".join(C.label(C.COMPLIANCE_TARGETS, t) for t in (project.compliance_targets or []))],
        ],
    )

    builder.heading("二、等保定级问卷答案")
    if survey and survey.answers_json:
        rows = [
            [ans.get("question_id", ""), ans.get("answer", "")]
            for ans in survey.answers_json
        ]
        builder.table(["问题编号", "答案描述"], rows)
    else:
        builder.para("尚未填写定级问卷。")

    builder.heading("三、定级结论")
    level = ctx.grading_level
    builder.para(
        f"经问卷评分判定,{project.name}的安全保护等级建议定为「{level or '未定级'}」。",
        bold=True,
    )

    builder.heading("四、判定理由")
    reason = survey.suggested_reason if survey else ""
    builder.para(reason or "（问卷结果未提供判定理由。）")

    builder.heading("五、判定依据")
    builder.para(
        "本建议书判定依据引用下列监管文件与标准(原文引用, 条款以现行有效版本为准):",
    )
    builder.bullets([
        "《信息安全技术 网络安全等级保护定级指南》(GB/T 22240-2020)——定级要素与定级流程;",
        "《金融行业网络安全等级保护实施指引》(JR/T 0071-2020)——金融行业等级保护实施与测评要求;",
        "《中华人民共和国网络安全法》第21条——国家网络安全等级保护制度;",
        "《信息安全等级保护管理办法》(公通字〔2007〕43号)第14、15条——三级系统年度测评与公安机关备案;",
    ])

    builder.heading("六、人工修正栏")
    adjusted = bool(survey and survey.final_level and survey.final_level != survey.suggested_level)
    if adjusted:
        builder.para(f"最终定级（已人工修正）：{survey.final_level}")
        builder.para(f"修正说明：{survey.manual_adjust_note or '无'}")
    else:
        builder.para(f"最终定级：{survey.final_level or '______________'}")
        builder.para("修正说明：______________________________________________________")
        builder.para("(本栏由定级复核人员在评审会上填写, 复核结论以人工修正值为准。)")

    builder.heading("七、安全中心复核意见")
    initiation = _gate_of(gates, "initiation")
    if initiation and (initiation.get("reviewer_opinion") or initiation.get("final_opinion")):
        builder.para(f"评审员意见：{initiation.get('reviewer_opinion') or '—'}")
        builder.para(f"负责人终审意见：{initiation.get('final_opinion') or '—'}")
        builder.para(
            f"门禁状态：{C.label(C.GATE_STATUSES, initiation.get('status', 'pending'))}; "
            f"交付物快照：{initiation.get('version_hash') or '—'}"
        )
    else:
        builder.para("复核结论：□ 同意定级结论    □ 调整为____级")
        builder.para("复核意见：______________________________________________________")
        builder.para("安全中心复核人签字：______________    日期：____________")
    return None


def _gate_of(gates: list | None, gate_type: str) -> dict | None:
    """从门禁快照列表中取指定类型门禁。"""
    for gate in (gates or []):
        if gate.get("gate_type") == gate_type:
            return gate
    return None


# ─────────────── 文档二: 需求规格说明书-安全需求章节 ───────────────


# 需求文档三组排序(改造点6): 监管报送类 → 等保条款类 → 通用安全类
DOC_GROUPS = ["监管报送类", "等保条款类", "通用安全类"]

# 等保相关监管文件关键词(判定"等保条款类")
_DJCP_KEYWORDS = ("等级保护", "等保", "GB/T 22239", "JR/T 0071", "JR/T 0068", "22239")


def compliance_basis(req: SecurityRequirement) -> list[str]:
    """需求 → 合规依据文案列表("《文件》第X条"), 取自 regulatory_ref。"""
    out = []
    for ref in (req.regulatory_ref or []):
        file_name = ref.get("file", "")
        if not file_name:
            continue
        clause = ref.get("clause", "")
        text = f"《{file_name}》{clause}" if clause else f"《{file_name}》"
        if ref.get("note"):
            text += f"({ref['note']})"
        out.append(text)
    return out or ["—"]


def compliance_group(req: SecurityRequirement) -> str:
    """三组归属: 监管报送类(regulatory_trigger) / 等保条款类 / 通用安全类。"""
    if req.category == C.label(C.TRIGGER_CATEGORY_LABELS, "regulatory_trigger"):
        return DOC_GROUPS[0]
    basis = "、".join(compliance_basis(req))
    if any(keyword in basis for keyword in _DJCP_KEYWORDS):
        return DOC_GROUPS[1]
    return DOC_GROUPS[2]


def group_by_compliance(requirements: list[SecurityRequirement]) -> list[tuple[str, list]]:
    """按"监管报送类 / 等保条款类 / 通用安全类"三组分组, 组序固定。"""
    groups: dict[str, list] = {name: [] for name in DOC_GROUPS}
    for req in requirements:
        groups.setdefault(compliance_group(req), []).append(req)
    return [(name, groups[name]) for name in DOC_GROUPS if groups.get(name)]


def group_by_asvs(requirements: list[SecurityRequirement]) -> list[tuple[str, list]]:
    """(保留)按 ASVS 章节分组, 供其他视图复用; 需求文档正文已改用三组排序。"""
    groups: dict[str, list] = {}
    order: list[str] = []
    for req in requirements:
        chapter = ""
        if req.asvs_ref:
            code = req.asvs_ref.split(".")[0].upper()
            if code in ASVS_CHAPTERS:
                chapter = f"{code} {ASVS_CHAPTERS[code]}"
        if not chapter:
            chapter = f"其他（{req.category}）"
        if chapter not in groups:
            groups[chapter] = []
            order.append(chapter)
        groups[chapter].append(req)

    def sort_key(chapter: str):
        num = "".join(filter(str.isdigit, chapter.split()[0]))
        return (1 if num else 0, int(num) if num else 0, chapter)

    return sorted(groups.items(), key=lambda pair: sort_key(pair[0]))


PRIORITY_ORDER = {p: i for i, p in enumerate(["critical", "high", "medium", "low"])}


def requirement_rows(items: list[SecurityRequirement]) -> tuple[list[list], set]:
    rows, red_cells = [], set()
    for r, req in enumerate(sorted(items, key=lambda x: (PRIORITY_ORDER.get(x.priority, 9), x.req_id))):
        if req.priority == "critical":
            red_cells.add((r, 0))
            red_cells.add((r, 1))
        rows.append([
            req.req_id,
            [req.title, req.description],
            C.label(C.REQUIREMENT_PRIORITY_LABELS, req.priority),
            compliance_basis(req),
            req.trigger_reason,
            req.acceptance_criteria,
        ])
    return rows, red_cells


def build_requirement_spec(ctx: RequirementContext, requirements: list, builder: DocBuilder,
                           gates: list | None = None) -> None:
    builder.heading("一、需求概览")
    by_category: dict[str, int] = {}
    for req in requirements:
        by_category[req.category] = by_category.get(req.category, 0) + 1
    stat_lines = [f"{cat}: {count} 条" for cat, count in by_category.items()]
    builder.para(f"本章共生成安全需求 {len(requirements)} 条，分类分布：" + "；".join(stat_lines) + "。")
    builder.para("全部需求均可追溯到来源输入(source_entity)，触发原因见各行『来源』列；"
                 "合规出处见『合规依据』列(条款号以合规部门确认为准)。")

    builder.heading("二、安全需求明细（监管报送类 / 等保条款类 / 通用安全类）")
    for group_name, items in group_by_compliance(requirements):
        builder.heading(f"{group_name}（{len(items)} 条）", level=2)
        rows, red_cells = requirement_rows(items)
        builder.table(
            ["需求编号", "需求描述", "优先级", "合规依据", "来源", "验收标准"],
            rows,
            red_cells=red_cells,
        )

    _append_permission_matrix(ctx, builder)
    _append_review_record(builder, gates)


def _append_review_record(builder: DocBuilder, gates: list | None) -> None:
    """评审记录页(改造点6): 评审节点/评审人/意见/签字栏/日期。"""
    builder.heading("评审记录")
    gates = gates or []
    if not gates:
        builder.para("本项目尚未进入评审流程, 以下记录由评审会后手工补签。")
    rows = []
    for gate in gates:
        rows.append([
            C.label(C.GATE_TYPES, gate.get("gate_type", "")),
            C.label(C.GATE_STATUSES, gate.get("status", "pending")),
            gate.get("submitter") or "—",
            gate.get("reviewer") or "",
            gate.get("reviewer_opinion") or "",
            (gate.get("final_reviewer") or "") + (f"/{gate.get('final_opinion')}" if gate.get("final_opinion") else ""),
        ])
    builder.table(
        ["评审节点", "状态", "提交人", "评审人", "评审意见", "终审/意见"],
        rows,
    )
    builder.para("签字确认(本人确认以上评审结论真实有效):")
    builder.para("项目经理签字: ______________    日期: ____________")
    builder.para("安全中心评审员签字: ______________    日期: ____________")
    builder.para("安全中心负责人签字: ______________    日期: ____________")
    builder.para(
        "注: 平台内电子留痕以「姓名+工号+时间戳+哈希」代替电子签章,"
        "哈希链校验详见《项目安全评审表》。"
    )


def _append_permission_matrix(ctx: RequirementContext, builder: DocBuilder) -> None:
    """附录: 角色×资源权限矩阵交叉表, 需审批操作加 * 标注。"""
    builder.heading("附录A 权限矩阵")
    resources = sorted(ctx.resources, key=lambda x: x.id)
    actions_order = list(C.PERMISSION_ACTIONS.keys())
    header = ["角色＼资源"] + [
        [res.name, f"({C.label(C.CRITICALITY_LEVELS, res.criticality)})"] for res in resources
    ]
    rows = []
    for role in sorted(ctx.roles, key=lambda x: x.id):
        row: list = [[
            role.name,
            f"{C.label(C.ROLE_TYPES, role.role_type)}·约{role.user_count_estimate}人",
        ]]
        for res in resources:
            entries = [
                e for e in ctx.permission_entries
                if e.role_id == role.id and e.resource_id == res.id
            ]
            marks = sorted(
                (
                    C.label(C.PERMISSION_ACTIONS, e.action) + ("*" if e.requires_approval else "")
                    for e in entries
                ),
                key=lambda m: next((i for i, a in enumerate(actions_order)
                                    if C.label(C.PERMISSION_ACTIONS, a) == m.rstrip("*")), 99),
            )
            row.append("、".join(marks) if marks else "—")
        rows.append(row)
    builder.table(header, rows)
    builder.para("注：操作后带 * 表示该操作需走审批流程；『—』表示未授予任何操作。", size=builder.style["fonts"]["table_size_pt"])


# ─────────────── 文档三: 总体设计说明书-安全设计章节 ───────────────


def build_design_baseline(ctx: RequirementContext, requirements: list,
                          builder: DocBuilder) -> None:
    project = ctx.project

    builder.heading("一、软件/框架版本清单")
    if ctx.components:
        rows = [
            [
                str(i + 1),
                C.label(C.SBOM_LAYERS, comp.layer),
                comp.name,
                comp.version,
                comp.purl or "",
                comp.license or "",
                C.label(C.SBOM_SOURCE_TYPES, comp.source_type),
            ]
            for i, comp in enumerate(sorted(ctx.components, key=lambda c: c.id))
        ]
        builder.table(["序号", "层级", "组件名称", "版本", "purl", "许可证", "录入来源"], rows)
    else:
        builder.para("暂无软件/框架清单。")

    builder.heading("二、数据字典（资产 → 数据表 → 字段, JR/T 0197 五级）")
    for asset in ctx.data_assets:
        storage = "、".join(C.label(C.STORAGE_ENVS, e) for e in (asset.storage_envs or []))
        builder.heading(asset.name, level=2)
        builder.table(
            ["资产属性", "取值"],
            [
                ["数据分类", C.label(C.DATA_ASSET_TYPES, asset.data_type)],
                ["安全分级(JR/T 0197-2020)", C.DATA_LEVEL_META.get(asset.classification, {}).get("label", asset.classification)],
                ["C3鉴别信息标签", _yesno(bool(asset.c3_tag))],
                *( [["迁移前分级留痕", asset.legacy_classification]] if getattr(asset, "legacy_classification", None) else [] ),
                ["是否个人信息", _yesno(asset.is_pii)],
                ["是否敏感个人信息", _yesno(asset.is_sensitive_pii)],
                ["存储位置", storage],
                ["涉及跨境传输", _yesno(asset.cross_border_transfer)],
            ],
        )
        for table in asset.tables:
            builder.para(f"数据表：{table.table_name}", bold=True)
            field_rows = [
                [
                    fld.field_name,
                    fld.field_type,
                    _yesno(fld.need_encrypt),
                    _yesno(fld.need_mask),
                    fld.mask_rule or "",
                ]
                for fld in table.fields
            ]
            builder.table(["字段名", "类型", "加密存储", "脱敏展示", "脱敏规则"], field_rows)

    builder.heading("三、API 接口安全属性表")
    if ctx.api_endpoints:
        rows = []
        for ep in ctx.api_endpoints:
            related = "、".join(ctx.sensitive_asset_names(ep.sensitive_asset_ids))
            rows.append([
                ep.name,
                f"{ep.method} {ep.path}",
                _yesno(ep.auth_required),
                _yesno(ep.public_exposed),
                related or "—",
                ep.rate_limit or "未配置",
            ])
        builder.table(["接口名称", "路径", "需认证", "公网暴露", "关联敏感数据资产", "限流"], rows)
    else:
        builder.para("暂无 API 接口清单。")

    builder.heading("四、基础设施资产清单")
    if ctx.infra_assets:
        rows = [
            [
                C.label(C.INFRA_ASSET_TYPES, a.asset_type),
                a.name,
                C.label(C.ENV_NAMES, a.env),
                a.ip or "",
                a.owner or "",
                _yesno(a.holds_sensitive),
            ]
            for a in ctx.infra_assets
        ]
        builder.table(["资产类型", "名称", "环境", "IP 地址", "负责人", "承载敏感数据"], rows)
    else:
        builder.para("暂无基础设施资产清单。")

    builder.heading("五、登录与密码策略设计说明")
    policy = effective_password_policy(ctx)
    builder.para(
        f"{project.name}的身份鉴别机制采用口令+动态因素组合。结合等保"
        f"{ctx.grading_level or '二级'}基线要求, 本系统登录口令策略设计如下:"
        f"口令最小长度不少于{policy['pwd_min_length']}位;"
        f"复杂度须至少包含{policy['pwd_complexity']}类字符(大写字母/小写字母/数字/特殊符号);"
        f"口令有效期{policy['pwd_valid_days']}天, 到期强制更换;"
        f"新口令不得与最近{policy['pwd_history_limit']}次历史口令重复;"
        f"连续{policy['lockout_threshold']}次验证失败锁定账户并通知安全审计后台;"
        f"会话空闲{policy['session_timeout_min']}分钟自动退出;"
        f"同一账号并发会话数上限{policy['concurrent_limit']}个。"
    )
    force_2fa_note = "系统强制启用双因素认证(口令+动态口令/短信验证码)。" if (
        ctx.auth_config and ctx.auth_config.force_2fa
    ) else "双因素认证按风险场景选择性启用。"
    builder.para(force_2fa_note)
    builder.table(
        ["策略项", "生效值"],
        [
            ["密码最小长度", f"{policy['pwd_min_length']} 位"],
            ["复杂度类别数", f"{policy['pwd_complexity']} 类"],
            ["有效期", f"{policy['pwd_valid_days']} 天"],
            ["历史口令限制", f"不得与前 {policy['pwd_history_limit']} 次重复"],
            ["错误锁定阈值", f"{policy['lockout_threshold']} 次"],
            ["会话超时", f"{policy['session_timeout_min']} 分钟"],
            ["并发会话上限", f"{policy['concurrent_limit']} 个"],
            ["强制双因素认证", _yesno(bool(ctx.auth_config and ctx.auth_config.force_2fa))],
        ],
    )

    builder.heading("六、认证方式设计说明")
    methods = ctx.auth_config.auth_methods if ctx.auth_config else []
    chosen = [(m, C.label(C.AUTH_METHODS, m)) for m in methods]
    if not chosen:
        builder.para("未选择认证方式。")
    for code, label_text in chosen:
        note = AUTH_METHOD_DESIGN_NOTES.get(code, "按行内统一安全规范实现。")
        builder.para(f"{label_text}：{note}")

    builder.heading("七、监管报送事项清单")
    regulatory_category = C.label(C.TRIGGER_CATEGORY_LABELS, "regulatory_trigger")
    reg_reqs = sorted(
        (r for r in requirements if r.category == regulatory_category),
        key=lambda r: r.req_id,
    )
    if not reg_reqs:
        builder.para("本项目未触发监管报送类事项。")
    else:
        rows = [
            [
                req.req_id,
                req.title,
                "、".join(compliance_basis(req)),
                _yesno(bool(req.reg_confirmed)),
                req.confirmed_by or "—",
                req.owner or "—",
            ]
            for req in reg_reqs
        ]
        builder.table(["编号", "报送事项", "监管依据", "已确认", "确认人", "责任人"], rows)
        builder.para("注: 报送事项须经项目经理逐条确认后方可通过立项门禁; 条款号以合规部门确认为准。")


# ─────────────── 文档四: SBOM 及漏洞清单 ───────────────


def build_sbom_vuln_report(
    ctx: RequirementContext,
    requirements: list,
    vulnerabilities: list,
    osv_summary: str,
    builder: DocBuilder,
) -> None:
    builder.heading("一、SBOM 总览")
    counts: dict[str, int] = {}
    for comp in ctx.components:
        counts[comp.layer] = counts.get(comp.layer, 0) + 1
    layer_stat = "；".join(f"{C.label(C.SBOM_LAYERS, k)} {v} 个" for k, v in counts.items())
    builder.para(
        f"共登记软件组件 {len(ctx.components)} 个（{layer_stat}）。"
        f"CycloneDX 1.5 格式清单随本文档一并交付, 供软件物料管理与漏洞持续监控导入使用。"
    )
    severity_counts: dict[str, int] = {}
    for v in vulnerabilities:
        severity_counts[v.severity] = severity_counts.get(v.severity, 0) + 1
    sev_stat = "；".join(
        f"{C.label(C.SEVERITY_LABELS, s)} {severity_counts[s]} 个"
        for s in ("critical", "high", "medium", "low") if severity_counts.get(s)
    )
    builder.para(f"漏洞查询结果：{osv_summary}。命中已知漏洞 {len(vulnerabilities)} 个（{sev_stat or '无'}）。")

    comp_name = {c.id: c for c in ctx.components}

    builder.heading("二、SBOM 组件清单")
    if ctx.components:
        rows = [
            [
                str(i + 1),
                C.label(C.SBOM_LAYERS, comp.layer),
                comp.name,
                comp.version,
                comp.license or "",
                C.label(C.SBOM_SOURCE_TYPES, comp.source_type),
                comp.purl or "",
            ]
            for i, comp in enumerate(sorted(ctx.components, key=lambda c: c.id))
        ]
        builder.table(["序号", "层级", "组件名称", "版本", "许可证", "来源", "purl"], rows)
    else:
        builder.para("暂无 SBOM 组件。")

    builder.heading("三、漏洞清单（高危/严重置顶标红）")
    if not vulnerabilities:
        builder.para("本次漏洞查询暂不可用或未发现已知漏洞, 结果不影响其余章节交付; 网络恢复后可重新执行查询。")
    else:
        high_levels = {"critical", "high"}
        red_rows = {
            (r, c) for r, v in enumerate(vulnerabilities)
            if v.severity in high_levels for c in range(7)
        }
        rows = [
            [
                v.cve_id,
                C.label(C.SEVERITY_LABELS, v.severity),
                f"{v.cvss_score:g}" if v.cvss_score is not None else "—",
                f"{c.name}@{c.version}" if (c := comp_name.get(v.component_id)) else "",
                v.affected_range or "—",
                v.fix_version or "官方暂未发布修复版",
                v.summary or "",
            ]
            for v in vulnerabilities
        ]
        builder.table(
            ["漏洞编号", "等级", "CVSS", "影响组件", "受影响范围", "修复版本", "简述"],
            rows,
            red_cells=red_rows,
        )

    builder.heading("四、漏洞整改联动需求")
    vuln_reqs = [req for req in requirements if req.source_entity_type == "sbom_component"]
    if not vuln_reqs:
        builder.para("当前无高危及以上漏洞, 未触发漏洞整改类安全需求。")
    else:
        def _target(req) -> str:
            comp = comp_name.get(req.source_entity_id)
            return f"{comp.name}@{comp.version}" if comp else ""

        rows = [
            [
                req.req_id,
                req.title,
                _target(req),
                C.label(C.REQUIREMENT_PRIORITY_LABELS, req.priority),
                req.description,
            ]
            for req in vuln_reqs
        ]
        builder.table(["需求编号", "整改需求", "目标组件", "优先级", "要求"], rows)


# ─────────────── 文档五: 项目安全评审表 ───────────────


def build_review_summary(
    ctx: RequirementContext,
    requirements: list,
    vulnerabilities: list,
    gates: list | None,
    builder: DocBuilder,
) -> None:
    """评审会会议材料: 门禁状态 + 需求覆盖统计 + 漏洞概况 + 遗留问题。"""
    project = ctx.project

    builder.heading("一、项目概况与门禁状态")
    gate_rows = []
    for gate in (gates or []):
        gate_rows.append([
            C.label(C.GATE_TYPES, gate.get("gate_type", "")),
            C.label(C.GATE_STATUSES, gate.get("status", "pending")),
            gate.get("submitter") or "—",
            gate.get("reviewer") or "—",
            gate.get("final_reviewer") or "—",
            (gate.get("version_hash") or "—")[:16] + ("…" if gate.get("version_hash") else ""),
        ])
    if gate_rows:
        builder.table(["评审节点", "状态", "提交人", "评审员", "终审人", "交付物快照(前16位)"], gate_rows)
    else:
        builder.para("各门禁尚未提交评审。")

    builder.heading("二、需求覆盖统计")
    by_category: dict[str, int] = {}
    by_priority: dict[str, int] = {}
    owner_missing = []
    for req in requirements:
        by_category[req.category] = by_category.get(req.category, 0) + 1
        by_priority[req.priority] = by_priority.get(req.priority, 0) + 1
        if req.priority == "critical" and not (req.owner or "").strip():
            owner_missing.append(req.req_id)
    builder.para(
        f"共 {len(requirements)} 条安全需求; 按优先级: " + "、".join(
            f"{C.label(C.REQUIREMENT_PRIORITY_LABELS, p)} {by_priority.get(p, 0)} 条"
            for p in ("critical", "high", "medium", "low")
        ) + "。"
    )
    builder.table(
        ["业务类目", "数量"],
        [[cat, str(count)] for cat, count in sorted(by_category.items())],
    )
    if owner_missing:
        builder.para(f"⚠ 以下 critical 需求未指定责任人: {'、'.join(owner_missing)}", bold=True)
    else:
        builder.para("全部 critical 需求均已指定责任人。")

    builder.heading("三、漏洞概况")
    severity_counts: dict[str, int] = {}
    for v in vulnerabilities:
        severity_counts[v.severity] = severity_counts.get(v.severity, 0) + 1
    builder.para(
        "漏洞分布: " + "、".join(
            f"{C.label(C.SEVERITY_LABELS, s)} {severity_counts.get(s, 0)} 个"
            for s in ("critical", "high", "medium", "low")
        ) + f"; 合计 {len(vulnerabilities)} 个。"
    )
    unfix = [v for v in vulnerabilities if v.severity in ("critical", "high")]
    if unfix:
        builder.para(
            "⚠ 高危及以上漏洞 "
            + "、".join(f"{v.cve_id}({C.label(C.SEVERITY_LABELS, v.severity)})" for v in unfix[:6])
            + (" 等" if len(unfix) > 6 else "")
            + " 需在评审会明确整改时限。", bold=True,
        )

    builder.heading("四、遗留问题清单")
    open_reqs = [
        r for r in requirements
        if r.status in ("open", "in_progress")
        and (r.priority in ("critical", "high") or r.category == C.label(C.TRIGGER_CATEGORY_LABELS, "regulatory_trigger"))
    ]
    if not open_reqs:
        builder.para("无高优先级遗留问题。")
    else:
        builder.table(
            ["编号", "遗留问题", "优先级", "责任人", "状态"],
            [
                [
                    r.req_id, r.title,
                    C.label(C.REQUIREMENT_PRIORITY_LABELS, r.priority),
                    r.owner or "未指定",
                    C.label(C.REQUIREMENT_STATUS, r.status),
                ]
                for r in sorted(open_reqs, key=lambda x: (PRIORITY_ORDER.get(x.priority, 9), x.req_id))
            ],
        )
    builder.para(
        "评审结论: □ 通过评审    □ 有条件通过(遗留问题限期整改)    □ 不通过",
    )
    builder.para("参会评审委员会签字栏: ____________________    日期: ____________")


# ────────────────────────── 对外主入口 ──────────────────────────


DOC_BUILDERS = {
    "grading": ("系统定级建议书", lambda ctx, reqs, vulns, summary, b, gates: build_grading_report(ctx, b, gates)),
    "requirement": ("需求规格说明书_安全需求章节", lambda ctx, reqs, vulns, summary, b, gates: build_requirement_spec(ctx, reqs, b, gates)),
    "design": ("总体设计说明书_安全设计章节", lambda ctx, reqs, vulns, summary, b, gates: build_design_baseline(ctx, reqs, b)),
    "sbom_vuln": ("SBOM及漏洞清单", lambda ctx, reqs, vulns, summary, b, gates: build_sbom_vuln_report(ctx, reqs, vulns, summary, b)),
    "review": ("项目安全评审表", lambda ctx, reqs, vulns, summary, b, gates: build_review_summary(ctx, reqs, vulns, gates, b)),
}


def generate_all_documents(
    ctx: RequirementContext,
    out_dir: str | Path,
    requirements: list[SecurityRequirement],
    vulnerabilities: list | None = None,
    osv_summary: str = "未执行",
    generated_at=None,
    gates: list | None = None,
) -> dict[str, Path]:
    """一次生成 5 份 Word 文档到 out_dir, 返回 {doc_type: 文件路径}。

    requirements 为规则引擎产物(须已实例化); vulnerabilities 为 OSV 同步后的记录列表;
    gates 为评审门禁快照(无评审数据时相关章节输出待补签栏)。
    """
    from datetime import datetime as _dt

    style = load_doc_style()
    out_dir = Path(out_dir)
    when = _fmt_date(generated_at or _dt.now())
    project = ctx.project
    signer = project.pm_name or ""

    paths: dict[str, Path] = {}
    for doc_type, (title, builder_fn) in DOC_BUILDERS.items():
        builder = DocBuilder(style)
        builder.cover(
            title=f"{project.name}\n{title}",
            project_name=project.name,
            project_code=project.code,
            generated_at_text=when,
            preparer=signer,
        )
        builder_fn(ctx, requirements, vulnerabilities or [], osv_summary, builder, gates)
        path = builder.save(out_dir / f"{project.code}_{title}.docx")
        paths[doc_type] = path
        logger.info("已生成 %s", path)
    return paths
